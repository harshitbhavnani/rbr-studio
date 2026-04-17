import os
import tempfile
from datetime import date
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# -----------------------------
# Formatting helpers
# -----------------------------
HEADER_FILL = "4472C4"  # Excel-like blue
HEADER_FONT = RGBColor(0xFF, 0xFF, 0xFF)
H1_COLOR = RGBColor(0x1F, 0x4E, 0x78)
H2_COLOR = RGBColor(0x2E, 0x5C, 0x8A)
SUBTLE = RGBColor(0x66, 0x66, 0x66)


def _set_run_font(run, size_pt: int = 11, bold: bool = False, color: Optional[RGBColor] = None):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a docx table cell; hex_color like '4472C4'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10, color: Optional[RGBColor] = None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    _set_run_font(run, size_pt=font_size, bold=bold, color=color)


def _safe_str(v: Any) -> str:
    if v is None:
        return ""
    return str(v)


def _fmt_money(x) -> str:
    try:
        return f"${float(x):,.0f}"
    except Exception:
        return _safe_str(x)


# Markdown-ish bullet parsing:
# - **Label:** content
# - content
import re
_BULLET_LABEL_RE = re.compile(r"^\-\s*\*\*([^*]+)\*\*:?\s*(.*)$")
_PLAIN_BULLET_RE = re.compile(r"^\-\s+(.*)$")
_PLAIN_LABEL_BULLET_RE = re.compile(r"^\-\s*([^:]+):\s*(.*)$")


def _add_bullets(doc: Document, text: str):
    if not text:
        p = doc.add_paragraph("Not provided.")
        p.paragraph_format.space_after = Pt(6)
        return

    for raw in str(text).splitlines():
        ln = (raw or "").strip()
        if not ln:
            continue

        m = _BULLET_LABEL_RE.match(ln)
        if m:
            label = (m.group(1) or "").strip().rstrip(":")
            content = (m.group(2) or "").strip()

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

            r1 = p.add_run(f"{label}: ")
            _set_run_font(r1, size_pt=11, bold=True)

            r2 = p.add_run(content)
            _set_run_font(r2, size_pt=11, bold=False)
            continue

        m_plain_label = _PLAIN_LABEL_BULLET_RE.match(ln)
        if m_plain_label:
            label = (m_plain_label.group(1) or "").strip().rstrip(":")
            content = (m_plain_label.group(2) or "").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{label}: ")
            _set_run_font(r1, size_pt=11, bold=True)
            r2 = p.add_run(content)
            _set_run_font(r2, size_pt=11, bold=False)
            continue

        m2 = _PLAIN_BULLET_RE.match(ln)
        if m2:
            content = (m2.group(1) or "").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(content)
            _set_run_font(r, size_pt=11, bold=False)
            continue

        p = doc.add_paragraph(ln)
        p.paragraph_format.space_after = Pt(6)


def _set_doc_styles(doc: Document):
    # Headings
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.bold = True
    h1.font.size = Pt(16)
    h1.font.color.rgb = H1_COLOR

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.bold = True
    h2.font.size = Pt(13)
    h2.font.color.rgb = H2_COLOR

    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Margins
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)


HeaderSpec = Tuple[str, str, Optional[float]]  # (key, label, width_in_inches_or_None)


def _add_table_from_records(
    doc: Document,
    records: Optional[List[Dict[str, Any]]],
    headers: Optional[Sequence[HeaderSpec]] = None,
    *,
    first_col_bold: bool = False,
    header_fill: str = HEADER_FILL,
    table_style: str = "Table Grid",
) -> None:
    """
    records: list of dicts
    headers: if provided, fixed columns in order: (key, label, width_in_inches or None)
             if not provided, uses keys from first record (in insertion order), width auto.
    """
    if not records:
        return

    if headers is None:
        keys = list(records[0].keys())
        headers = [(k, k, None) for k in keys]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = table_style
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for j, (k, label, width_in) in enumerate(headers):
        _set_cell_text(hdr_cells[j], label, bold=True, font_size=10, color=HEADER_FONT)
        _shade_cell(hdr_cells[j], header_fill)
        if width_in:
            hdr_cells[j].width = Inches(width_in)

    # Data rows
    for rec in records:
        row_cells = table.add_row().cells
        for j, (k, label, width_in) in enumerate(headers):
            val = rec.get(k, "")
            _set_cell_text(
                row_cells[j],
                _safe_str(val),
                bold=(first_col_bold and j == 0),
                font_size=10,
            )
            if width_in:
                row_cells[j].width = Inches(width_in)


def _add_kv_table(doc: Document, kv: Dict[str, Any]):
    """
    Renders a 2-row table with keys as headers and values in the next row.
    """
    if not kv:
        return
    keys = list(kv.keys())
    table = doc.add_table(rows=2, cols=len(keys))
    table.style = "Table Grid"
    table.autofit = True

    # header row
    for j, k in enumerate(keys):
        _set_cell_text(table.cell(0, j), _safe_str(k), bold=True, font_size=10, color=HEADER_FONT)
        _shade_cell(table.cell(0, j), HEADER_FILL)

    # value row
    for j, k in enumerate(keys):
        _set_cell_text(table.cell(1, j), _safe_str(kv.get(k, "")), bold=False, font_size=10)


def _add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size_pt=10, bold=True)
    p.paragraph_format.space_after = Pt(4)


def _add_title_page(doc: Document, borrower_name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Commercial Real Estate")
    _set_run_font(r, size_pt=20, bold=True, color=H1_COLOR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Credit Analysis Report")
    _set_run_font(r, size_pt=20, bold=True, color=H1_COLOR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(borrower_name or "Borrower")
    _set_run_font(r, size_pt=16, bold=False, color=H2_COLOR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date.today().strftime("%B %d, %Y"))
    _set_run_font(r, size_pt=12, bold=False, color=SUBTLE)


def _add_client_info_section(doc: Document, borrower_name, client_number, branch_number, relationship_manager, loan_exposure, deposit_relationship, tier_level):
    doc.add_heading("Client Information", level=1)

    t = doc.add_table(rows=6, cols=3)
    t.style = "Table Grid"
    t.autofit = True

    _set_cell_text(t.cell(0, 0), "Client Group", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(0, 1), "Branch Number", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(0, 2), "Relationship Manager", bold=True, color=HEADER_FONT)
    for c in range(3):
        _shade_cell(t.cell(0, c), HEADER_FILL)

    _set_cell_text(t.cell(1, 0), borrower_name)
    _set_cell_text(t.cell(1, 1), branch_number)
    _set_cell_text(t.cell(1, 2), relationship_manager)

    _set_cell_text(t.cell(2, 0), "Client Number", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(2, 1), "Loan Exposure", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(2, 2), "Deposit Relationship", bold=True, color=HEADER_FONT)
    for c in range(3):
        _shade_cell(t.cell(2, c), HEADER_FILL)

    _set_cell_text(t.cell(3, 0), client_number)
    _set_cell_text(t.cell(3, 1), _fmt_money(loan_exposure))
    _set_cell_text(t.cell(3, 2), _fmt_money(deposit_relationship))

    _set_cell_text(t.cell(4, 0), "Tier Level", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(4, 1), "Report Date", bold=True, color=HEADER_FONT)
    _set_cell_text(t.cell(4, 2), "", bold=True, color=HEADER_FONT)
    for c in range(3):
        _shade_cell(t.cell(4, c), HEADER_FILL)

    _set_cell_text(t.cell(5, 0), tier_level)
    _set_cell_text(t.cell(5, 1), date.today().isoformat())
    _set_cell_text(t.cell(5, 2), "")
    doc.add_paragraph("")


def _add_commentary_block(doc: Document, rb: str, ai: str):
    if rb and rb.strip():
        doc.add_heading("Rule-Based Commentary", level=2)
        _add_bullets(doc, rb)
    if ai and ai.strip():
        doc.add_heading("AI-Generated Commentary", level=2)
        _add_bullets(doc, ai)


def _add_analysis_sections(
    doc: Document,
    *,
    loan_title: Optional[str] = None,
    loan_meta: Optional[Dict[str, Any]] = None,
    primary_repayment="",
    rent_roll_rule_based="",
    rent_roll_genai="",
    stress_test_rule_based="",
    stress_test_genai="",
    dscr_calculations="",
    rollover_rule_based="",
    rollover_genai="",
    construction_rule_based="",
    construction_genai="",
    rent_roll_table=None,
    lease_exp_table=None,
    vac_occ_table=None,
    stress_entry_table=None,
    stress_cap_table=None,
    stress_ir_table=None,
    stress_vac_table=None,
    stress_noi_table=None,
    rollover_table=None,
    constr_noi_scenarios=None,
    constr_cap_ltv_table=None,
    constr_ir_dscr_table=None,
    constr_sale_table=None,
    constr_sizing=None,
):
    if loan_title:
        doc.add_heading(loan_title, level=1)
    if loan_meta:
        doc.add_heading("Loan Snapshot", level=2)
        _add_kv_table(doc, loan_meta)
        doc.add_paragraph("")

    doc.add_heading("Primary Source of Repayment Analysis", level=1)
    if primary_repayment and primary_repayment.strip():
        _add_bullets(doc, primary_repayment)
    else:
        doc.add_paragraph("Primary source of repayment to be determined.")
    doc.add_page_break()

    doc.add_heading("Rent Roll Analysis", level=1)
    if rent_roll_table:
        doc.add_heading("Rent Roll Table", level=2)
        has_lease_start = isinstance(rent_roll_table, list) and rent_roll_table and ("lease_start" in rent_roll_table[0])
        rr_headers: List[HeaderSpec] = [("suite", "Suite", None), ("tenant", "Tenant", None), ("sf", "SF", None), ("pct_sf", "%SF", None)]
        if has_lease_start:
            rr_headers.append(("lease_start", "Lease Start", None))
        rr_headers += [
            ("lease_end", "Lease End", None),
            ("rem_term", "Rem Trm", None),
            ("base_rent", "Base Rent", None),
            ("cams", "CAMs", None),
            ("annual", "Annual", None),
            ("pct_rent", "%Rent", None),
        ]
        _add_table_from_records(doc, rent_roll_table, rr_headers, first_col_bold=True)
        doc.add_paragraph("")
    if lease_exp_table:
        doc.add_heading("Lease Expiration Summary", level=2)
        _add_table_from_records(doc, lease_exp_table, headers=[("Metric", "Metric", None), ("Year 1", "Year 1", None), ("Year 2", "Year 2", None), ("Year 3", "Year 3", None)], first_col_bold=True)
        doc.add_paragraph("")
    if vac_occ_table:
        doc.add_heading("Vacancy / Occupancy", level=2)
        _add_table_from_records(doc, vac_occ_table, headers=[("Status", "Status", None), ("sq_ft", "Sq Ft", None), ("pct", "%", None)], first_col_bold=True)
        doc.add_paragraph("")
    _add_commentary_block(doc, rent_roll_rule_based, rent_roll_genai)
    doc.add_page_break()

    doc.add_heading("Investor CRE Stress Test", level=1)
    if stress_entry_table:
        doc.add_heading("Entry Table", level=2)
        _add_table_from_records(doc, stress_entry_table, headers=[("Item", "Item", None), ("Source", "Source", None), ("Terms", "Terms", None)], first_col_bold=False)
        doc.add_paragraph("")
    doc.add_heading("Stress Tables", level=2)
    if stress_cap_table:
        _add_caption(doc, "Cap Rate Effect (LTV)")
        _add_table_from_records(doc, stress_cap_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    if stress_ir_table:
        _add_caption(doc, "Interest Rate Effect (DSCR)")
        _add_table_from_records(doc, stress_ir_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    if stress_vac_table:
        _add_caption(doc, "Vacancy Rate Effect (DSCR)")
        _add_table_from_records(doc, stress_vac_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    if stress_noi_table:
        _add_caption(doc, "NOI Change Effect (DSCR)")
        _add_table_from_records(doc, stress_noi_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    _add_commentary_block(doc, stress_test_rule_based, stress_test_genai)
    doc.add_page_break()

    doc.add_heading("Debt Coverage Ratio Calculations", level=1)
    if dscr_calculations and dscr_calculations.strip():
        _add_bullets(doc, dscr_calculations)
    else:
        doc.add_paragraph("DSCR calculations to be added.")
    doc.add_page_break()

    doc.add_heading("Rollover Risk Analysis", level=1)
    if rollover_table:
        doc.add_heading("Lease Rollover Risk Table", level=2)
        _add_table_from_records(doc, rollover_table, headers=[("Row", "", None), ("Year 1", "Year 1", None), ("Year 2", "Year 2", None), ("Year 3", "Year 3", None)], first_col_bold=True)
        doc.add_paragraph("")
    _add_commentary_block(doc, rollover_rule_based, rollover_genai)
    doc.add_page_break()

    doc.add_heading("Construction & Bridge Loan Analysis", level=1)
    if constr_noi_scenarios:
        doc.add_heading("NOI Scenarios", level=2)
        if isinstance(constr_noi_scenarios, dict):
            _add_kv_table(doc, constr_noi_scenarios)
            doc.add_paragraph("")
    if constr_sizing:
        doc.add_heading("Quick Takeout Sizing", level=2)
        sizing_records = [{
            "Max Takeout @ LTV": constr_sizing.get("max_ltv", ""),
            "Max Takeout @ DSCR": constr_sizing.get("max_dscr", ""),
            "Bridge Commitment": constr_sizing.get("bridge", ""),
        }]
        _add_table_from_records(doc, sizing_records, headers=[("Max Takeout @ LTV", "Max Takeout @ LTV", None), ("Max Takeout @ DSCR", "Max Takeout @ DSCR", None), ("Bridge Commitment", "Bridge Commitment", None)], first_col_bold=False)
        doc.add_paragraph("")
    if constr_cap_ltv_table:
        doc.add_heading("Cap Rate → Max Takeout (LTV)", level=2)
        _add_table_from_records(doc, constr_cap_ltv_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    if constr_ir_dscr_table:
        doc.add_heading("Interest Rate → Max Takeout (DSCR)", level=2)
        _add_table_from_records(doc, constr_ir_dscr_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    if constr_sale_table:
        doc.add_heading("Cap Rate → Sale Prices", level=2)
        _add_table_from_records(doc, constr_sale_table, headers=None, first_col_bold=True)
        doc.add_paragraph("")
    _add_commentary_block(doc, construction_rule_based, construction_genai)


# -----------------------------
# Main export function
# -----------------------------
def create_credit_analysis_report(
    borrower_name="",
    client_number="",
    branch_number="",
    relationship_manager="",
    loan_exposure=0.0,
    deposit_relationship=0.0,
    tier_level="",
    exec_summary="",
    primary_repayment="",
    rent_roll_rule_based="",
    rent_roll_genai="",
    stress_test_rule_based="",
    stress_test_genai="",
    dscr_calculations="",
    rollover_rule_based="",
    rollover_genai="",
    construction_rule_based="",
    construction_genai="",
    rent_roll_table=None,
    lease_exp_table=None,
    vac_occ_table=None,
    stress_entry_table=None,
    stress_cap_table=None,
    stress_ir_table=None,
    stress_vac_table=None,
    stress_noi_table=None,
    rollover_table=None,
    constr_noi_scenarios=None,
    constr_cap_ltv_table=None,
    constr_ir_dscr_table=None,
    constr_sale_table=None,
    constr_sizing=None,
    loans_data: Optional[List[Dict[str, Any]]] = None,
    output_path: Optional[str] = None,
):
    """
    Generates a Word document to output_path and returns the file path.

    When loans_data is provided, one combined relationship document is created with a section for each loan.
    """
    if output_path is None:
        output_path = os.path.join(tempfile.gettempdir(), "credit_analysis_report.docx")

    out_dir = os.path.dirname(output_path) or "."
    os.makedirs(out_dir, exist_ok=True)

    doc = Document()
    _set_doc_styles(doc)
    _add_title_page(doc, borrower_name)
    doc.add_page_break()
    _add_client_info_section(doc, borrower_name, client_number, branch_number, relationship_manager, loan_exposure, deposit_relationship, tier_level)
    doc.add_heading("Executive Relationship Summary", level=1)
    doc.add_paragraph((exec_summary or "No executive summary provided.").strip())

    if loans_data:
        for idx, loan_data in enumerate(loans_data):
            doc.add_page_break()
            _add_analysis_sections(doc, **loan_data)
            if idx < len(loans_data) - 1:
                doc.add_page_break()
    else:
        doc.add_page_break()
        _add_analysis_sections(
            doc,
            primary_repayment=primary_repayment,
            rent_roll_rule_based=rent_roll_rule_based,
            rent_roll_genai=rent_roll_genai,
            stress_test_rule_based=stress_test_rule_based,
            stress_test_genai=stress_test_genai,
            dscr_calculations=dscr_calculations,
            rollover_rule_based=rollover_rule_based,
            rollover_genai=rollover_genai,
            construction_rule_based=construction_rule_based,
            construction_genai=construction_genai,
            rent_roll_table=rent_roll_table,
            lease_exp_table=lease_exp_table,
            vac_occ_table=vac_occ_table,
            stress_entry_table=stress_entry_table,
            stress_cap_table=stress_cap_table,
            stress_ir_table=stress_ir_table,
            stress_vac_table=stress_vac_table,
            stress_noi_table=stress_noi_table,
            rollover_table=rollover_table,
            constr_noi_scenarios=constr_noi_scenarios,
            constr_cap_ltv_table=constr_cap_ltv_table,
            constr_ir_dscr_table=constr_ir_dscr_table,
            constr_sale_table=constr_sale_table,
            constr_sizing=constr_sizing,
        )

    doc.save(output_path)
    return output_path